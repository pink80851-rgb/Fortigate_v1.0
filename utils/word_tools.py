import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def apply_kaishu_style(paragraph):
    """強制將段落內的所有文字設為 12 點標楷體"""
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.name = '標楷體'
        # 針對中文字型的特殊處理
        r = run._element.get_or_add_rPr()
        rFonts = r.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '標楷體')

def handle_word_file(request_obj, template_path, storage_folder):
    try:
        if not os.path.exists(storage_folder):
            os.makedirs(storage_folder)

        doc = Document(template_path)

        # 1. 準備替換資料 (根據你的需求填寫)
        ext_access_text = f"{request_obj.ext_ip}:{request_obj.ext_port}" if request_obj.ext_ip else "220.128.111.105"
        
        replacements = {
            "{{ dept }}": str(request_obj.dept or ""),
            "{{ applicant }}": str(request_obj.applicant or ""),
            "{{ reason }}": str(request_obj.reason or ""),
            "{{ floor }}": str(request_obj.floor or ""),
            "{{ int_ip }}": str(request_obj.int_ip or "未分配"),
            "{{ ext_access }}": ext_access_text,
            "{{ create_date }}": request_obj.created_at.strftime('%Y-%m-%d') if request_obj.created_at else "",
            "□網路連接設定": "■網路連接設定",
            "□白名單申請": "■白名單申請"
        }

        # 2. 開始在表格與段落中尋找並替換
        # 處理普通段落
        for paragraph in doc.paragraphs:
            for target, value in replacements.items():
                if target in paragraph.text:
                    # 使用這種方式替換可以保留部分格式，但我們最後會統一重設標楷體
                    paragraph.text = paragraph.text.replace(target, value)
                    apply_kaishu_style(paragraph)

        # 處理表格內部的段落 (你的範本看起來很多都在表格裡)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for target, value in replacements.items():
                            if target in paragraph.text:
                                paragraph.text = paragraph.text.replace(target, value)
                                apply_kaishu_style(paragraph)

        # 3. 儲存檔案
        save_name = f"{request_obj.ticket_id}.docx"
        final_path = os.path.join(storage_folder, save_name)
        doc.save(final_path)
        return final_path

    except Exception as e:
        print(f"❌ Word 產出異常: {str(e)}")
        return None